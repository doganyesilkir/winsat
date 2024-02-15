import subprocess
import xml.etree.ElementTree as ET
import glob
import os
import ctypes
import sys
import time
from docx import Document



kayityolu=Document()

if __name__ == "__main__":
        if ctypes.windll.shell32.IsUserAnAdmin() == 0:
            ctypes.windll.shell32.ShellExecuteW(None, "runas", sys.executable, " ".join(sys.argv), None, 1)
            
klasor_yolu="C:\\Windows\\Performance\\WinSat\\DataStore"

try:
            dosya_listesi = os.listdir(klasor_yolu)

            for dosya_adi in dosya_listesi:
                dosya_yolu = os.path.join(klasor_yolu, dosya_adi)
            
                if os.path.isfile(dosya_yolu):
                    try:
                        os.remove(dosya_yolu)
                        print(f"{dosya_adi} başarıyla silindi.")
                    except Exception as e:
                        print(f"{dosya_adi} silinirken bir hata oluştu: {e}")
except Exception as e:
    print(f"Klasör içindeki dosyalar silinirken bir hata oluştu: {e}")

os.system("cls")
print("Harici Terminal Kapandığında Masaüstünüzde Winsat Sonuçları Adında Bir Word Belgesi Oluşacaktır..")
komut = "winsat formal"
subprocess.run(komut, shell=True)


os.system("cls")


def find_file_with_partial_name(partial_name, starting_path="."):
    pattern = os.path.join(starting_path, f"*{partial_name}*")
    matches = glob.glob(pattern)
    return matches if matches else None

def xml_dosyasini_parse_et(xml_dosya_yolu):
    tree = ET.parse(xml_dosya_yolu)
    root = tree.getroot()

    for degerler in root.findall('WinSPR'):
        systemscore = degerler.find('SystemScore').text
        memoryscore = degerler.find('MemoryScore').text
        cpuscore = degerler.find('CpuScore').text
        cpusubaggscore = degerler.find('CPUSubAggScore').text
        videoencodescore = degerler.find('VideoEncodeScore').text
        graphicsscore = degerler.find('GraphicsScore').text
        dx9subscore = degerler.find('Dx9SubScore').text
        dx10subscore = degerler.find('Dx10SubScore').text
        gamingscore = degerler.find('GamingScore').text
        diskscore = degerler.find('DiskScore').text
        kayityolu.add_paragraph(f"Dosya: {xml_dosya_yolu}")
        kayityolu.add_paragraph(f"Sistem Puanı: {systemscore} \nRam Puanı: {memoryscore} \nİşlemci Puanı: {cpuscore} \nİşlemci Genel Performans Puanı: {cpusubaggscore} \nVideo Encode Puanı: {videoencodescore} \nGrafik Puanı: {graphicsscore} \nDX9 Puanı: {dx9subscore} \nDX10 Puanı: {dx10subscore} \nOyun Puanı: {gamingscore} \nDisk Puanı: {diskscore}")
    kayityolu.save("C:\\Users\Pc\OneDrive\Masaüstü\Winsat Sonuçları.docx")

kismi_isim="Formal.Assessment"

dosya_yollari = find_file_with_partial_name(kismi_isim, klasor_yolu)
if dosya_yollari:
    kayityolu.add_paragraph("Sistem Hız Testi Sonuçları:")
    for dosya_yolu in dosya_yollari:
        xml_dosyasini_parse_et(dosya_yolu)
        sys.exit()
else:
    print(f"{kismi_isim} içeren dosya bulunamadı.")
